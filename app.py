import json
import os
import shutil
import sqlite3
import tempfile
import uuid
import re
from datetime import datetime
import csv
PIL_AVAILABLE = False
TESSERACT_AVAILABLE = False
WINRT_AVAILABLE = False
import os as _os
try:
    from PIL import Image
    PIL_AVAILABLE = True
except Exception:
    Image = None

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except Exception:
    pytesseract = None


def _configure_tesseract_cmd():
    """tesseract sering terinstal tapi tidak ada di PATH (khususnya di Windows).
    Arahkan pytesseract ke binary-nya agar OCR benar-benar berjalan."""
    if not TESSERACT_AVAILABLE:
        return
    if shutil.which("tesseract"):
        return  # sudah ada di PATH
    candidates = [
        _os.environ.get("TESSERACT_CMD"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.join(_os.environ.get("LOCALAPPDATA", ""), "Programs", "Tesseract-OCR", "tesseract.exe"),
    ]
    for cand in candidates:
        if cand and os.path.exists(cand):
            pytesseract.pytesseract.tesseract_cmd = cand
            break


_configure_tesseract_cmd()

try:
    import asyncio
    from winrt.windows.graphics.imaging import BitmapDecoder, BitmapPixelFormat, BitmapAlphaMode, SoftwareBitmap
    from winrt.windows.media.ocr import OcrEngine
    from winrt.windows.storage.streams import InMemoryRandomAccessStream, DataWriter
    WINRT_AVAILABLE = True
except Exception:
    asyncio = None
    BitmapDecoder = None
    BitmapPixelFormat = None
    BitmapAlphaMode = None
    SoftwareBitmap = None
    OcrEngine = None
    InMemoryRandomAccessStream = None
    DataWriter = None

_USE_GOOGLE_VISION = False
try:
    from google.cloud import vision
    # enable if credentials env var is set
    if _os.environ.get('GOOGLE_APPLICATION_CREDENTIALS'):
        _USE_GOOGLE_VISION = True
except Exception:
    _USE_GOOGLE_VISION = False

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.table import _Cell
from flask import Flask, flash, redirect, render_template, request, send_from_directory, url_for, jsonify
from werkzeug.utils import secure_filename

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "akuisisi.db")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
GENERATED_DIR = os.path.join(BASE_DIR, "generated")
TEMPLATE_PATH = os.path.join(BASE_DIR, "template_laporan.docx")
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
OUTPUT_FONT_NAME = "Arial"
OUTPUT_FONT_SIZE_PT = 11
TITLE_FONT_SIZE_PT = 14
PHOTO_COLS_PER_ROW = 2
# Larger default photo width for clearer images in generated reports
PHOTO_WIDTH_INCHES = 1.8
PHOTO_WIDTH_MM = 45
# Saat True, hasil generate akan mempertahankan format template apa adanya.
USE_STRICT_TEMPLATE_FORMAT = True

# Teks petunjuk di template Word -> kolom database (None = isi khusus / gambar)
HINT_FIELD_MAP = {
    "Diisi dengan tanggal pelaporan": "tanggal_laporan",
    "Diisi dengan tanggal dan jam": "tanggal_serah_terima",
    "Diisi denga nama tempat pelaksanaan amuses": "nama_instansi",
    "Diisi dengan alat forensik digital yang digunakan": "keterangan_barang",
    "Diisi dengan metode akuisisi": "prosedur_akuisisi",
    "Diisi nama file hasil akuisisi": "nomor_laporan",
    "Diisi dengan nilai hash hasil akuisisi": "keterangan_barang",
    "Diisi status akuisisi": "status",
}

# Informasi Perangkat (Hasil Akuisisi) — petunjuk di template Word tabel Informasi Perangkat
PERANGKAT_INFO_FIELDS = [
    ("os", "Operasi Sistem (OS)", "Diisi dengan nama tempat pelaksanaan akuisisi"),
    ("os_version", "Versi OS", "Diisi dengan versi OS"),
    ("imei", "IMEI", "Diisi dengan IMEI"),
    ("imsi", "IMSI", "Diisi dengan IMSI"),
    ("iccid", "ICCID", "Diisi dengan ICCID"),
    ("msisdn", "MSISDN", "Diisi dengan MSISDN"),
    ("serial_number", "Serial Number", "Diisi Serial Number Jika Ada"),
]
PERANGKAT_INFO_HINT_MAP = {hint: key for key, _label, hint in PERANGKAT_INFO_FIELDS}
PERANGKAT_INFO_LABEL_MAP = {label.lower(): key for key, label, _hint in PERANGKAT_INFO_FIELDS}

MAIN_PHOTO_HINT = "Diisi foto barang bukti elektronik"
SCREENSHOT_HINTS = [
    "Diisi dengan Gambar Tangkapan layar nilai hash",
    "Diisi dengan Gambar Tangkapan layar ekstraksi hasil akuisisi",
    "Diisi dengan Gambar Tangkapan layar indeksing hasil akuisisi",
]

# Poin H: Data yang Ditemukan (kolom form -> kunci JSON -> teks petunjuk di template Word)
OCR_DATA_FIELDS = [
    ("chats", "Chats", "Diisi dengan jumlah chat"),
    ("contacts", "Contacts", "Diisi dengan jumlah kontak"),
    ("installed_applications", "Installed Applications", "Diisi dengan jumlah aplikasi terinstal"),
    ("instant_messages", "Instant Messages", "Diisi dengan jumlah Instant Messages"),
    ("user_accounts", "User Accounts", "Diisi dengan jumlah akun pengguna"),
    ("timeline", "Timeline", "Diisi dengan jumlah timeline"),
    ("applications", "Applications", "Diisi dengan jumlah aplikasi"),
    ("archives", "Archives", "Diisi dengan jumlah archives"),
    ("databases", "Databases", "Diisi dengan jumlah databases"),
    ("images", "Images", "Diisi dengan jumlah gambar"),
    ("text", "Text", "Diisi dengan jumlah teks"),
]
OCR_DATA_HINT_MAP = {hint: key for key, _label, hint in OCR_DATA_FIELDS}

# Data predefined untuk Surat Dinas dan Surat Tugas (dapat dimodifikasi)
PREDEFINED_SURATS = {
    "surat_dinas": [
        {
            "nomor": "R-PD.04.03.4B.10.25.614",
            "tanggal": "2026-01-20",
            "perihal": "perihal Permohonan Bantuan Analisis Digital Forensik"
        }
    ],
    "surat_tugas": [
        {
            "nomor": "R-PD.04.03.18B.01.26.45",
            "tanggal": "2026-01-20",
            "perihal": "perihal Kegiatan Pemeriksaan Forensik Digital terhadap Barang Bukti Elektronik Kasus"
        }
    ]
}

app = Flask(__name__)
app.secret_key = "ganti-secret-key-aman"


def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def load_instansi():
    csv_path = os.path.join(BASE_DIR, 'data', 'instansi.csv')
    instansi = []
    if not os.path.exists(csv_path):
        return instansi
    with open(csv_path, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            name = row.get('Nama UPT') or row.get('Nama') or ''
            alamat = row.get('Alamat') or ''
            instansi.append({'name': name.strip(), 'alamat': alamat.strip()})
    # Debug: print how many entries were loaded and which file path was used
    try:
        print(f"DEBUG: load_instansi loaded {len(instansi)} entries from {csv_path}")
    except Exception:
        pass
    return instansi


def init_db():
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)
    ensure_default_template()
    conn = get_conn()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS laporan_akuisisi (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nomor_laporan TEXT NOT NULL UNIQUE,
            tanggal_laporan TEXT NOT NULL,
            nama_instansi TEXT,
            alamat_instansi TEXT,
            tanggal_serah_terima TEXT,
            komoditi TEXT,
            jenis_pelanggaran TEXT,
            jenis_barang_bukti TEXT,
            merek_model TEXT,
            imei TEXT,
            keterangan_barang TEXT,
            ocr_raw_text TEXT,
            ocr_fields_json TEXT,
            ocr_summary_text TEXT,
            dasar_items TEXT,
            foto_paths TEXT,
            hash_foto_paths TEXT,
            ekstraksi_foto_paths TEXT,
            indexing_foto_paths TEXT,
            waktu_dimulai TEXT,
            waktu_selesai TEXT,
            tempat_pelaksanaan TEXT,
            perangkat_forensik_digital TEXT,
            metode_akuisisi TEXT,
            status_akuisisi TEXT,
            nama_file_hash TEXT,
            nilai_hash TEXT,
            maksud_tujuan TEXT,
            prosedur_akuisisi TEXT,
            kesimpulan TEXT,
            status TEXT NOT NULL DEFAULT 'draft',
            created_at TEXT NOT NULL
        )
        """
    )
    
    # Tambahkan kolom baru jika belum ada (untuk database existing)
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN hash_foto_paths TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN ekstraksi_foto_paths TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN indexing_foto_paths TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN waktu_dimulai TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN waktu_selesai TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN tempat_pelaksanaan TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN perangkat_forensik_digital TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN metode_akuisisi TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN status_akuisisi TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN nama_file_hash TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN nilai_hash TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN nomor_surat_dinas TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN tanggal_surat_dinas TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN ocr_raw_text TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN ocr_fields_json TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN ocr_summary_text TEXT")
    except:
        pass
    try:
        conn.execute("ALTER TABLE laporan_akuisisi ADD COLUMN perangkat_info_json TEXT")
    except:
        pass
    
    conn.commit()
    conn.close()

    # ensure counters table for atomic nomor generation
    conn = get_conn()
    conn.execute("CREATE TABLE IF NOT EXISTS counters (name TEXT PRIMARY KEY, value INTEGER)")
    try:
        conn.execute("INSERT OR IGNORE INTO counters (name, value) VALUES ('nomor_laporan', 0)")
        conn.commit()
    except Exception:
        pass
    conn.close()


def ensure_default_template():
    if os.path.exists(TEMPLATE_PATH):
        return

    doc = Document()
    doc.add_heading("Laporan Hasil Akuisisi Barang Bukti Elektronik {{merek_model}}", 1)
    doc.add_paragraph("Nomor Laporan: {{nomor_laporan}}")
    doc.add_paragraph("Tanggal Laporan: {{tanggal_laporan}}")

    doc.add_heading("A. Dasar", level=2)
    doc.add_paragraph("{{dasar_items}}")

    doc.add_heading("B. Identitas Pemangku Kepentingan", level=2)
    doc.add_paragraph("Nama Instansi: {{nama_instansi}}")
    doc.add_paragraph("Alamat Instansi: {{alamat_instansi}}")
    doc.add_paragraph("Nomor Surat Dinas: {{nomor_surat_dinas}}")
    doc.add_paragraph("Tanggal Surat Dinas: {{tanggal_surat_dinas}}")
    doc.add_paragraph("Tanggal Serah Terima: {{tanggal_serah_terima}}")

    doc.add_heading("C. Jenis Kasus", level=2)
    doc.add_paragraph("Komoditi: {{komoditi}}")
    doc.add_paragraph("Jenis Pelanggaran: {{jenis_pelanggaran}}")

    doc.add_heading("D. Identitas Barang Bukti Elektronik", level=2)
    doc.add_paragraph("Jenis Barang Bukti: {{jenis_barang_bukti}}")
    doc.add_paragraph("Merek/Model: {{merek_model}}")
    doc.add_paragraph("IMEI: {{imei}}")
    doc.add_paragraph("Keterangan Barang: {{keterangan_barang}}")
    doc.add_paragraph("Daftar Foto Upload: {{foto_list}}")

    doc.add_heading("E. Maksud dan Tujuan", level=2)
    doc.add_paragraph("{{maksud_tujuan}}")

    doc.add_heading("F. Prosedur Akuisisi", level=2)
    doc.add_paragraph("{{prosedur_akuisisi}}")

    doc.add_heading("G. Pelaksanaan Akuisisi", level=2)
    table = doc.add_table(rows=8, cols=2)
    table.cell(0, 0).text = "Waktu dimulai"
    table.cell(0, 1).text = "{{waktu_dimulai}}"
    table.cell(1, 0).text = "Waktu selesai"
    table.cell(1, 1).text = "{{waktu_selesai}}"
    table.cell(2, 0).text = "Tempat"
    table.cell(2, 1).text = "{{tempat_pelaksanaan}}"
    table.cell(3, 0).text = "Perangkat forensik digital yang digunakan"
    table.cell(3, 1).text = "{{perangkat_forensik_digital}}"
    table.cell(4, 0).text = "Metode Akuisisi"
    table.cell(4, 1).text = "{{metode_akuisisi}}"
    table.cell(5, 0).text = "Status"
    table.cell(5, 1).text = "{{status_akuisisi}}"
    table.cell(6, 0).text = "Nama file"
    table.cell(6, 1).text = "{{nama_file_hash}}"
    table.cell(7, 0).text = "Nilai Hash"
    table.cell(7, 1).text = "{{nilai_hash}}"

    doc.add_heading("H. Data yang Ditemukan dari OCR", level=2)
    doc.add_paragraph("{{ocr_summary_text}}")

    doc.add_heading("I. Kesimpulan", level=2)
    doc.add_paragraph("{{kesimpulan}}")

    doc.save(TEMPLATE_PATH)


def _apply_font_to_runs(runs, size_pt=OUTPUT_FONT_SIZE_PT, force_bold=None):
    for run in runs:
        run.font.name = OUTPUT_FONT_NAME
        run.font.size = Pt(size_pt)
        if force_bold is not None:
            run.bold = force_bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), OUTPUT_FONT_NAME)


def normalize_docx_font(docx_path):
    # Keep exact Word formatting from user template when strict mode is enabled.
    if USE_STRICT_TEMPLATE_FORMAT:
        return

    # Force output text style so generated values are not inheriting unwanted style.
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name:
            _apply_font_to_runs(paragraph.runs, size_pt=TITLE_FONT_SIZE_PT, force_bold=True)
        else:
            _apply_font_to_runs(paragraph.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_font_to_runs(paragraph.runs)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _apply_font_to_runs(paragraph.runs)
        for paragraph in section.footer.paragraphs:
            _apply_font_to_runs(paragraph.runs)
    doc.save(docx_path)


def allowed_file(filename):
    if "." not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in ALLOWED_EXTENSIONS


def format_tanggal(value):
    if not value:
        return "-"
    raw = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            dt = datetime.strptime(raw, fmt)
            bulan = [
                "Januari",
                "Februari",
                "Maret",
                "April",
                "Mei",
                "Juni",
                "Juli",
                "Agustus",
                "September",
                "Oktober",
                "November",
                "Desember",
            ]
            return f"{dt.day} {bulan[dt.month - 1]} {dt.year}"
        except ValueError:
            continue
    return raw


def field_value(row, field_name):
    if not field_name:
        return "-"
    value = row[field_name]
    if value is None or str(value).strip() == "":
        return "-"
    if field_name in ("tanggal_laporan", "tanggal_serah_terima"):
        return format_tanggal(value)
    return str(value).strip()


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def merge_jinja_runs_in_paragraph(paragraph):
    runs = paragraph.runs
    if not runs:
        return
    combined = "".join(run.text or "" for run in runs)
    if "{{" not in combined and "{%" not in combined and "}}" not in combined:
        return
    runs[0].text = combined
    for run in runs[1:]:
        run.text = ""


def repair_template_jinja_tags(docx_path):
    doc = Document(docx_path)
    for paragraph in iter_all_paragraphs(doc):
        merge_jinja_runs_in_paragraph(paragraph)
    doc.save(docx_path)


def normalize_pelaksanaan_template(docx_path):
    doc = Document(docx_path)

    if len(doc.tables) >= 4:
        tbl = doc.tables[3]
        tbl._element.getparent().remove(tbl._element)

    anchor = None
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip().upper()
        if text == "PELAKSANAAN AKUISISI" or text == "G. PELAKSANAAN AKUISISI":
            anchor = paragraph
            break

    if anchor is None:
        doc.save(docx_path)
        return

    items = [
        ("Waktu dimulai", "{{waktu_dimulai}}"),
        ("Waktu selesai", "{{waktu_selesai}}"),
        ("Tempat", "{{tempat_pelaksanaan}}"),
        ("Perangkat forensik digital yang digunakan", "{{perangkat_forensik_digital}}"),
        ("Metode Akuisisi", "{{metode_akuisisi}}"),
        ("Status", "{{status_akuisisi}}"),
        ("Nama file", "{{nama_file_hash}}"),
        ("Nilai Hash", "{{nilai_hash}}"),
    ]

    def insert_after(paragraph, text=""):
        new_element = OxmlElement("w:p")
        paragraph._element.addnext(new_element)
        new_paragraph = Paragraph(new_element, paragraph._parent)
        if text:
            new_paragraph.add_run(text)
        return new_paragraph

    current = anchor
    for idx, (label, placeholder) in enumerate(items, start=1):
        current = insert_after(current, f"{idx}. {label}: {placeholder}")
        current.style = doc.styles["Normal"]
        _set_numbered_paragraph_format(current, left_inch=0.5, hanging_inch=0.25)

    doc.save(docx_path)


def set_paragraph_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def clear_paragraph(paragraph):
    p_element = paragraph._element
    for child in list(p_element):
        if child.tag != qn("w:pPr"):
            p_element.remove(child)


def replace_hint_text(paragraph, hint, replacement):
    if hint not in (paragraph.text or ""):
        return False
    new_text = paragraph.text.replace(hint, replacement, 1)
    set_paragraph_text(paragraph, new_text)
    return True


def _set_table_borderless(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def _paragraph_parent_cell(paragraph):
    element = paragraph._element
    while element is not None:
        if element.tag == qn("w:tc"):
            return _Cell(element, paragraph.part)
        element = element.getparent()
    return None


def insert_images_in_paragraph(paragraph, image_paths, width_inches=PHOTO_WIDTH_INCHES, cols=PHOTO_COLS_PER_ROW, is_landscape=False):
    valid_paths = []
    for filename in image_paths:
        full_path = os.path.join(UPLOAD_DIR, filename)
        if os.path.isfile(full_path):
            valid_paths.append(full_path)

    if not valid_paths:
        set_paragraph_text(paragraph, "-")
        return

    cell = _paragraph_parent_cell(paragraph)
    clear_paragraph(paragraph)

    if cell is None:
        col_idx = 0
        for idx, path in enumerate(valid_paths):
            if idx > 0 and col_idx == 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run()
            if col_idx == 1:
                run.add_text("\t")
            run.add_picture(path, width=Inches(width_inches))
            col_idx = (col_idx + 1) % cols
        return

    cols_used = 1 if is_landscape else cols
    rows_needed = max(1, (len(valid_paths) + cols_used - 1) // cols_used)
    table = cell.add_table(rows=rows_needed, cols=cols_used)
    _set_table_borderless(table)

    # For landscape (indexing) screenshots use near-full width; otherwise use
    # the configured photo width.
    landscape_width = Inches(6.5) if is_landscape else Inches(width_inches)
    for idx, path in enumerate(valid_paths):
        row_idx, col_idx = divmod(idx, cols_used)
        img_cell = table.rows[row_idx].cells[col_idx]
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(path, width=landscape_width)


def fill_template_hints(doc, row, foto_bb_paths, hash_paths, indexing_paths):
    tanggal_jam_count = 0
    screenshot_idx = 0
    processed_paragraphs = set()

    for paragraph in iter_all_paragraphs(doc):
        text = (paragraph.text or "").strip()
        paragraph_id = id(paragraph._element)

        if text == MAIN_PHOTO_HINT:
            if paragraph_id in processed_paragraphs:
                continue
            processed_paragraphs.add(paragraph_id)
            if foto_bb_paths:
                insert_images_in_paragraph(paragraph, foto_bb_paths)
                merek = field_value(row, "merek_model")
                imei = field_value(row, "imei")
                caption = f"Gambar 1. Foto Barang Bukti Elektronik merek {merek}, IMEI {imei}"
                caption_par = _insert_caption_below(paragraph, caption)
                try:
                    processed_paragraphs.add(id(caption_par._element))
                except Exception:
                    pass
            else:
                set_paragraph_text(paragraph, "-")
            continue

        if text in SCREENSHOT_HINTS:
            if paragraph_id in processed_paragraphs:
                continue
            processed_paragraphs.add(paragraph_id)

            if screenshot_idx == 0:
                # Nilai Hash
                if hash_paths:
                    insert_images_in_paragraph(paragraph, hash_paths, width_inches=2.8, cols=1)
                    cap = _insert_caption_below(paragraph, "Gambar 2. Foto Tangkapan layar nilai hash")
                    try:
                        processed_paragraphs.add(id(cap._element))
                    except Exception:
                        pass
                else:
                    replace_hint_text(paragraph, text, "-")
            elif screenshot_idx == 1:
                # Ekstraksi (bisa multiple)
                ekstraksi_foto_paths = json.loads(row["ekstraksi_foto_paths"] or "[]")
                if ekstraksi_foto_paths:
                    insert_images_in_paragraph(paragraph, ekstraksi_foto_paths, width_inches=2.8, cols=1)
                    cap = _insert_caption_below(paragraph, "Gambar 3. Foto Tangkapan layar hasil ekstraksi")
                    try:
                        processed_paragraphs.add(id(cap._element))
                    except Exception:
                        pass
                else:
                    replace_hint_text(paragraph, text, "-")
            elif screenshot_idx == 2:
                # Indexing (landscape/full width)
                if indexing_paths:
                    insert_images_in_paragraph(paragraph, indexing_paths, is_landscape=True)
                    cap = _insert_caption_below(paragraph, "Gambar 4. Foto Tangkapan layar hasil indexing")
                    try:
                        processed_paragraphs.add(id(cap._element))
                    except Exception:
                        pass
                else:
                    replace_hint_text(paragraph, text, "-")

            screenshot_idx += 1
            continue

        if text in HINT_FIELD_MAP:
            field_name = HINT_FIELD_MAP[text]
            if text == "Diisi dengan tanggal dan jam":
                value = field_value(row, "tanggal_serah_terima" if tanggal_jam_count == 0 else "tanggal_laporan")
                tanggal_jam_count += 1
            else:
                value = field_value(row, field_name)
            replace_hint_text(paragraph, text, value)
            continue

        if text in OCR_DATA_HINT_MAP:
            ocr_fields = _load_ocr_fields(row)
            value = ocr_fields.get(OCR_DATA_HINT_MAP[text], "") or "-"
            replace_hint_text(paragraph, text, value)
            continue

        if text in PERANGKAT_INFO_HINT_MAP:
            perangkat_info = _load_perangkat_info(row)
            value = perangkat_info.get(PERANGKAT_INFO_HINT_MAP[text], "") or "-"
            replace_hint_text(paragraph, text, value)
            continue

        if text.startswith("Diisi dengan jumlah") or text.startswith("Diisi dengan nama petugas") or text.startswith(
            "Diisi dengan nama jabatan petugas"
        ) or text.startswith("Diisi dengan tandatangan elektronik"):
            replace_hint_text(paragraph, text, "-")

        if "Gambar. Foto Barang Bukti Elektronik" in text or "Gambar 1. Foto Barang Bukti Elektronik" in text:
            if paragraph_id in processed_paragraphs:
                continue
            processed_paragraphs.add(paragraph_id)
            _delete_paragraph(paragraph)


def build_render_context(doc, row, dasar_items, foto_bb_paths):
    dasar_lines = []
    for idx, item in enumerate(dasar_items, start=1):
        isi = _normalize_numbered_item_text(item.get('isi', ''))
        if isi:
            dasar_lines.append(f"{idx}. {isi}")
    dasar_text = "\n".join(dasar_lines)
    foto_text = ", ".join(foto_bb_paths) if foto_bb_paths else "-"
    kesimpulan = row["kesimpulan"] or "-"

    foto_images = []
    for filename in foto_bb_paths:
        full_path = os.path.join(UPLOAD_DIR, filename)
        if os.path.isfile(full_path):
            foto_images.append(InlineImage(doc, full_path, width=Inches(PHOTO_WIDTH_INCHES)))

    context = {
        "nomor_laporan": row["nomor_laporan"],
        "tanggal_laporan": format_tanggal(row["tanggal_laporan"]),
        "nama_instansi": row["nama_instansi"] or "-",
        "alamat_instansi": row["alamat_instansi"] or "-",
        "nomor_surat_dinas": row["nomor_surat_dinas"] or "-",
        "tanggal_surat_dinas": format_tanggal(row["tanggal_surat_dinas"]),
        "tanggal_serah_terima": format_tanggal(row["tanggal_serah_terima"]),
        "komoditi": row["komoditi"] or "-",
        "jenis_pelanggaran": row["jenis_pelanggaran"] or "-",
        "jenis_barang_bukti": row["jenis_barang_bukti"] or "-",
        "merek_model": row["merek_model"] or "-",
        # Backwards-compatibility: some templates use uppercase or alternate keys.
        # Provide an uppercase variant for templates that expect all-caps placeholders.
        "MEREK_MODEL": (row["merek_model"] or "-").upper(),
        "MerekModel": row["merek_model"] or "-",
        "merek_tipe_model": row["merek_model"] or "-",
        "merek_tipe": row["merek_model"] or "-",
        "tipe_model": row["merek_model"] or "-",
        "imei": row["imei"] or "-",
        "keterangan_barang": row["keterangan_barang"] or "-",
        "ocr_raw_text": row["ocr_raw_text"] or "-",
        "ocr_fields_json": row["ocr_fields_json"] or "-",
        "ocr_summary_text": row["ocr_summary_text"] or "-",
        "dasar_items": dasar_text or "-",
        "maksud_tujuan": row["maksud_tujuan"] or "-",
        "prosedur_akuisisi": _normalize_prosedur_text_for_render(row["prosedur_akuisisi"] or "") or "-",
        "kesimpulan": kesimpulan,
        "Kesimpulan": kesimpulan,
        "waktu_dimulai": row["waktu_dimulai"] or "-",
        "waktu_selesai": row["waktu_selesai"] or "-",
        "tempat_pelaksanaan": row["tempat_pelaksanaan"] or "-",
        "perangkat_forensik_digital": row["perangkat_forensik_digital"] or "-",
        "metode_akuisisi": row["metode_akuisisi"] or "-",
        "status_akuisisi": row["status_akuisisi"] or "-",
        "nama_file_hash": row["nama_file_hash"] or "-",
        "nilai_hash": row["nilai_hash"] or "-",
        "foto_list": foto_text,
        "foto_images": foto_images,
        "foto1": foto_images[0] if len(foto_images) > 0 else "-",
        "foto2": foto_images[1] if len(foto_images) > 1 else "-",
        "foto3": foto_images[2] if len(foto_images) > 2 else "-",
        "Nomor_Laporan": row["nomor_laporan"],
        "NomorLaporan": row["nomor_laporan"],
    }
    return context


def _delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def _insert_paragraph_after(paragraph, text="", style=None):
    new_paragraph_element = OxmlElement("w:p")
    paragraph._element.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _insert_paragraph_before(paragraph, text="", style=None):
    new_paragraph_element = OxmlElement("w:p")
    paragraph._element.addprevious(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _set_numbered_paragraph_format(paragraph, left_inch=0.5, hanging_inch=0.5):
    # Make numbering visually consistent with manual reports: a slightly larger
    # left indent and a clear hanging indent for wrapped lines.
    paragraph.paragraph_format.left_indent = Inches(left_inch)
    paragraph.paragraph_format.first_line_indent = Inches(-hanging_inch)
    
    # Clear existing tab stops and add new one at the left indent position
    try:
        paragraph.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    
    try:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(left_inch))
    except Exception:
        pass
    
    # Keep numbered lines tight so 3 and 4 do not split with extra vertical gap
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0


def _insert_numbered_item_after(paragraph, number, body, number_width_inch=0.25):
    new_paragraph = _insert_paragraph_after(paragraph, "")
    # Set format BEFORE setting text for better formatting application
    _set_numbered_paragraph_format(new_paragraph)
    set_paragraph_text(new_paragraph, f"{number}.\t{body}")
    return new_paragraph


def _split_intro_and_numbered_items(text):
    intro_lines = []
    numbered_lines = []
    seen_numbered = False

    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.match(r"^\d+\.\s*", line):
            seen_numbered = True
            numbered_lines.append(line)
        elif seen_numbered:
            if numbered_lines:
                numbered_lines[-1] = f"{numbered_lines[-1]} {line}".strip()
        else:
            intro_lines.append(line)

    intro_text = " ".join(intro_lines).strip()
    numbered_items = _compose_numbered_items("\n".join(numbered_lines)) if numbered_lines else []
    return intro_text, numbered_items


def _insert_caption_below(paragraph, caption_text):
    cell = _paragraph_parent_cell(paragraph)
    if cell is None:
        caption_paragraph = _insert_paragraph_after(paragraph, "")
    else:
        caption_paragraph = cell.add_paragraph()
    set_paragraph_text(caption_paragraph, caption_text)
    _set_tight_paragraph_format(caption_paragraph)
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.size = Pt(10)
    return caption_paragraph


def _set_tight_paragraph_format(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0


# Panjang hex baku tiap algoritma (jumlah karakter hex).
_HASH_LEN = {"md5": 32, "sha1": 40, "sha256": 64}
# Pola label per baris, toleran salah-baca OCR (mis. "SHA1"->"SHAL",
# "MD5"->"MDS"/"Mos"). Diurut spesifik→umum agar SHA256 dicoba sebelum SHA1.
_HASH_LINE_LABELS = [
    ("sha256", r"SHA[\s\-]*2[\s\-]*[5S][\s\-]*6"),
    ("sha1",   r"SHA[\s\-]*[1lI](?!\d)|SHA[L1lI]"),
    ("md5",    r"M[D0O][\s\-]*[5S]|MOS"),
]
# Label algoritma LAIN (bukan target) untuk menghindari salah atribusi pada
# fallback berbasis panjang, karena panjang hex-nya bisa sama dengan target.
_OTHER_HASH_LABELS = re.compile(
    r"MD\s*4|MD\s*2|SHA\s*384|SHA\s*512|RIPEMD\s*160|PANAMA|TIGER|ADLER\s*32|CRC\s*32|eDonkey",
    re.IGNORECASE,
)


def parse_hash_values(text):
    """Ekstrak nilai MD5/SHA1/SHA256 dari teks hasil OCR foto hash.

    Pendekatan per baris: deteksi label (toleran salah-baca OCR) lalu ambil
    deretan hex sesudahnya. Untuk baris tanpa label, fallback klasifikasi
    deret hex 'bersih' berdasarkan panjang unik (32/40/64).

    Mengembalikan (hashes, warnings): hashes = {'md5','sha1','sha256'} (hanya
    yang ditemukan, lowercase); warnings = peringatan panjang tak sesuai.
    Catatan: OCR hex panjang rawan salah baca — hasil selalu perlu diverifikasi."""
    text = text or ""
    hashes = {}
    warnings = []

    # 1) Per baris: cari label, ambil hex sesudahnya (buang semua non-hex).
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        for key, pattern in _HASH_LINE_LABELS:
            if key in hashes:
                continue
            m = re.search(pattern, line, re.IGNORECASE)
            if not m:
                continue
            expected = _HASH_LEN[key]
            hexonly = re.sub(r"[^0-9a-fA-F]", "", line[m.end():]).lower()
            if len(hexonly) < 8:
                continue
            value = hexonly[:expected]
            hashes[key] = value
            if len(value) != expected:
                warnings.append(
                    f"{key.upper()} terbaca {len(value)} karakter (seharusnya {expected}) - periksa manual."
                )
            break

    # 2) Fallback berbasis panjang HANYA bila tak ada satu pun label target
    #    terbaca (mis. dump hex polos). Banyak algoritma berbagi panjang
    #    (MD4/MD2/eDonkey=32, RIPEMD160=40, PANAMA=64), jadi lewati baris yang
    #    jelas berlabel algoritma LAIN agar tak salah atribusi.
    if not hashes:
        for raw in text.splitlines():
            line = raw.strip()
            if not line or _OTHER_HASH_LABELS.search(line):
                continue
            for token in re.findall(r"(?<![0-9a-fA-F])[0-9a-fA-F]{32,64}(?![0-9a-fA-F])", line):
                token = token.lower()
                for key, expected in _HASH_LEN.items():
                    if key not in hashes and len(token) == expected:
                        hashes[key] = token
                        break

    return hashes, warnings


def _ocr_hash_image(path):
    """OCR khusus gambar nilai hash. Grayscale + upscale tanpa threshold keras
    memberi pembacaan hex terbaik pada screenshot seperti HashCalc (jauh lebih
    baik daripada _ocr_with_pytesseract_variants yang early-return pada gambar
    mentah). Memilih hasil dengan karakter hex terbanyak."""
    if not PIL_AVAILABLE or not TESSERACT_AVAILABLE:
        return ""
    try:
        img = Image.open(path)
        if img.mode not in ("L", "RGB"):
            img = img.convert("RGB")
        w, h = img.size
        scale = 3 if max(w, h) < 2000 else 2
        gray = img.convert("L").resize((w * scale, h * scale))
    except Exception:
        return ""
    best, best_score = "", -1
    for config in ("--oem 3 --psm 6", "--oem 3 --psm 4", "--oem 3 --psm 11"):
        try:
            t = pytesseract.image_to_string(gray, lang="eng", config=config)
        except Exception:
            continue
        score = len(re.findall(r"[0-9a-fA-F]", t or ""))
        if score > best_score:
            best, best_score = (t or ""), score
    return best


def format_nilai_hash_text(hashes):
    """Format dict hash menjadi teks bernomor sesuai placeholder field
    'Nilai Hash': '1. MD5 : ...\\n2. SHA1 : ...\\n3. SHA256 : ...'."""
    order = [("md5", "MD5"), ("sha1", "SHA1"), ("sha256", "SHA256")]
    lines = []
    idx = 1
    for key, label in order:
        value = (hashes or {}).get(key)
        if value:
            lines.append(f"{idx}. {label} : {value}")
            idx += 1
    return "\n".join(lines)


def _parse_ocr_summary_stats(text):
    text = (text or "").strip()
    if not text:
        return {}
    if text.startswith("{"):
        try:
            data = json.loads(text)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}

    stats = {}
    label_to_key = {label.lower(): key for key, label, _hint in OCR_DATA_FIELDS}
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = re.match(r"^([^:]+):\s*(.+)$", line)
        if not match:
            continue
        label = match.group(1).strip().lower()
        value = match.group(2).strip()
        for known_label, key in label_to_key.items():
            if known_label == label or known_label in label or label in known_label:
                stats[key] = value
                break
    return stats


def _load_ocr_fields(row):
    row = dict(row) if not isinstance(row, dict) else row
    raw_json = (row.get("ocr_fields_json") or "").strip()
    if raw_json:
        try:
            data = json.loads(raw_json)
            if isinstance(data, dict):
                return data
        except Exception:
            pass
    return _parse_ocr_summary_stats(row.get("ocr_summary_text") or "")


def build_ocr_data_from_form(form):
    fields = {}
    for key, _label, _hint in OCR_DATA_FIELDS:
        value = (form.get(f"ocr_{key}", "") or "").strip()
        if value:
            fields[key] = value
    return fields


def build_ocr_summary_text_from_fields(fields):
    lines = []
    for key, label, _hint in OCR_DATA_FIELDS:
        value = (fields.get(key) or "").strip()
        if value:
            lines.append(f"{label}: {value}")
    return "\n".join(lines)


def build_pelaksanaan_akuisisi_values(form):
    return {
        "waktu_dimulai": form.get("waktu_dimulai", "").strip(),
        "waktu_selesai": form.get("waktu_selesai", "").strip(),
        "tempat_pelaksanaan": form.get("tempat_pelaksanaan", "").strip(),
        "perangkat_forensik_digital": form.get("perangkat_forensik_digital", "").strip(),
        "metode_akuisisi": form.get("metode_akuisisi", "").strip(),
        "status_akuisisi": form.get("status_akuisisi", "").strip(),
        "nama_file_hash": form.get("nama_file_hash", "").strip(),
        "nilai_hash": form.get("nilai_hash", "").strip(),
    }


def build_perangkat_info_from_form(form):
    fields = {}
    for key, _label, _hint in PERANGKAT_INFO_FIELDS:
        value = (form.get(f"perangkat_{key}", "") or "").strip()
        if value:
            fields[key] = value
    return fields


def _load_perangkat_info(row):
    row = dict(row) if not isinstance(row, dict) else row
    raw_json = (row.get("perangkat_info_json") or "").strip()
    info = {}
    if raw_json:
        try:
            data = json.loads(raw_json)
            if isinstance(data, dict):
                info = data
        except Exception:
            pass
    if not info.get("imei") and (row.get("imei") or "").strip():
        info["imei"] = (row.get("imei") or "").strip()
    return info


def _perangkat_info_display_value(info, key):
    value = (info.get(key) or "").strip()
    return value if value else "-"


def _fill_perangkat_info_table(doc, row):
    info = _load_perangkat_info(row)
    if not info:
        return

    target_table = None
    for table in doc.tables:
        try:
            if table.rows and "Informasi Perangkat" in (table.rows[0].cells[0].text or ""):
                target_table = table
                break
        except Exception:
            continue
    if target_table is None:
        return

    for row_idx in range(1, len(target_table.rows)):
        try:
            label = (target_table.rows[row_idx].cells[0].text or "").strip().lower()
            key = PERANGKAT_INFO_LABEL_MAP.get(label)
            if not key:
                continue
            value = _perangkat_info_display_value(info, key)
            if value == "-":
                continue
            target_table.rows[row_idx].cells[1].text = value
        except Exception:
            continue


def _fill_ocr_summary_table(doc, row):
    stats = _load_ocr_fields(row)
    if not stats:
        return

    target_table = None
    for table in doc.tables:
        try:
            if table.rows and len(table.rows[0].cells) >= 1 and "Data yang Ditemukan" in (table.rows[0].cells[0].text or ""):
                target_table = table
                break
        except Exception:
            continue

    if target_table is None:
        return

    row_map = [
        (1, "chats"),
        (2, "contacts"),
        (3, "installed_applications"),
        (4, "instant_messages"),
        (5, "user_accounts"),
        (6, "timeline"),
        (7, "applications"),
        (8, "archives"),
        (9, "databases"),
        (10, "images"),
        (11, "text"),
    ]

    for row_idx, key in row_map:
        if row_idx >= len(target_table.rows):
            continue
        value = stats.get(key)
        if value is None or str(value).strip() == "":
            continue
        try:
            target_table.rows[row_idx].cells[1].text = str(value)
        except Exception:
            pass


def _clean_numbered_lines(text):
    lines = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^\d+\.\s*", "", line)
        lines.append(line)
    return lines


def _normalize_numbered_item_text(text):
    return re.sub(r"\s+", " ", (text or "")).strip()


def _compose_numbered_items(text):
    items = []
    current = None
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        m = re.match(r"^(\d+)\.\s*(.*)$", line)
        if m:
            if current:
                items.append(_normalize_numbered_item_text(current))
            current = f"{m.group(1)}. {_normalize_numbered_item_text(m.group(2))}".strip()
            continue
        if current:
            current = f"{current} {_normalize_numbered_item_text(line)}".strip()
        else:
            current = line
    if current:
        items.append(_normalize_numbered_item_text(current))
    return items


def _normalize_prosedur_text_for_render(text):
    items = _compose_numbered_items(text)
    if items:
        return "\n".join(items)
    return (text or "").strip()


def build_prosedur_akuisisi(prosedur_option=None, prosedur_text=""):
    base = [
        "1. ISO/IEC 27037:2012 tentang Pedoman Identifikasi, Pengumpulan, Akuisisi dan Preservasi Bukti Digital;",
        "2. SOP Mikro No. POM-05.03/CFM.01/SOP.01/IK.63.08 tentang Penanganan Barang Bukti Elektronik;",
        "3. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.16 tentang Instruksi Kerja Akuisisi Handphone;",
    ]

    option = (prosedur_option or "").strip().lower()
    if option == "oxygen":
        base.append(
            "4. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.19 tentang Instruksi Kerja Penggunaan Oxygen Forensic Extractor."
        )
    elif option == "cellebrite":
        base.extend([
            "4. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.17 tentang Instruksi Kerja Penggunaan Cellebrite UFED.",
            "5. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.23 tentang Instruksi Kerja Penggunaan Aplikasi Analisis Physical Analyzer.",
        ])
    elif option == "both":
        base.extend([
            "4. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.17 tentang Instruksi Kerja Penggunaan Cellebrite UFED.",
            "5. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.23 tentang Instruksi Kerja Penggunaan Aplikasi Analisis Physical Analyzer.",
            "6. Instruksi Kerja No. POM-05.03/CFM.01/SOP.01/IK.63.19 tentang Instruksi Kerja Penggunaan Oxygen Forensic Extractor.",
        ])

    if option:
        return "\n".join(base)
    return (prosedur_text or "").strip()


def format_numbered_sections(doc, dasar_items, prosedur_text):
    paragraphs = doc.paragraphs

    def find_heading_index(keyword):
        for idx, paragraph in enumerate(paragraphs):
            if keyword.lower() in paragraph.text.lower():
                return idx
        return None

    dasar_idx = find_heading_index("A. Dasar")
    if dasar_idx is not None and dasar_idx + 1 < len(paragraphs):
        proc_end = dasar_idx + 1
        # Remove the placeholder/content paragraph after the heading.
        _delete_paragraph(paragraphs[dasar_idx + 1])
        anchor = paragraphs[dasar_idx]
        current = anchor
        for idx, item in enumerate(dasar_items, start=1):
            body = _normalize_numbered_item_text(item.get('isi', ''))
            if not body:
                continue
            current = _insert_numbered_item_after(current, idx, body)
            current.style = doc.styles["Normal"]
            _set_numbered_paragraph_format(current, left_inch=0.5, hanging_inch=0.5)

    proc_idx = find_heading_index("F. Prosedur Akuisisi")
    if proc_idx is not None and proc_idx + 1 < len(paragraphs):
        proc_end = len(paragraphs)
        for idx in range(proc_idx + 1, len(paragraphs)):
            text = (paragraphs[idx].text or "").strip()
            if re.match(r"^[A-Z]\.", text):
                proc_end = idx
                break
        for idx in range(proc_end - 1, proc_idx, -1):
            _delete_paragraph(paragraphs[idx])
        anchor = paragraphs[proc_idx]
        current = _insert_paragraph_after(anchor, "Pelaksanaan akuisisi menggunakan:")
        current.style = doc.styles["Normal"]
        _set_tight_paragraph_format(current)
        # Normalize stored procedure text, split into numbered items, and
        # insert each item as a separate numbered paragraph (same style as dasar).
        normalized = _normalize_prosedur_text_for_render(prosedur_text)
        proc_items = _compose_numbered_items(normalized)
        for idx, item_text in enumerate(proc_items, start=1):
            # strip any leading numeric prefix from parsed item and re-insert
            body = re.sub(r"^\d+\.\s*", "", item_text).strip()
            current = _insert_numbered_item_after(current, idx, body)
            current.style = doc.styles["Normal"]
            _set_numbered_paragraph_format(current, left_inch=0.5, hanging_inch=0.5)


def align_dasar_indent_to_prosedur(doc):
    """Samakan indentasi daftar bernomor pada bagian DASAR dengan bagian
    PROSEDUR AKUISISI. Keduanya dirender sebagai satu paragraf placeholder,
    namun template memberi format berbeda (DASAR: first-line indent, PROSEDUR:
    left indent blok). Salin format paragraf PROSEDUR ke paragraf DASAR agar
    penomorannya sejajar."""
    paras = doc.paragraphs

    def numbered_para_after(heading_text):
        heading_text = heading_text.lower()
        seen = False
        for paragraph in paras:
            text = (paragraph.text or "").strip()
            if not seen:
                if text.lower() == heading_text:
                    seen = True
                continue
            if re.match(r"^\d+\.", text):
                return paragraph
            # Berhenti bila sudah masuk heading bagian berikutnya (huruf kapital).
            if text and text.isupper() and len(text) > 3:
                return None
        return None

    dasar_p = numbered_para_after("dasar")
    prosedur_p = numbered_para_after("prosedur akuisisi")
    if dasar_p is None or prosedur_p is None:
        return

    src = prosedur_p.paragraph_format
    dst = dasar_p.paragraph_format
    dst.left_indent = src.left_indent
    dst.first_line_indent = src.first_line_indent
    dst.space_before = src.space_before
    dst.space_after = src.space_after
    try:
        dasar_p.style = prosedur_p.style
    except Exception:
        pass


def format_kesimpulan_section(doc, kesimpulan_text):
    paragraphs = doc.paragraphs

    def find_heading_index(keyword):
        for idx, paragraph in enumerate(paragraphs):
            if keyword.lower() in (paragraph.text or "").lower():
                return idx
        return None

    kesimpulan_idx = find_heading_index("I. Kesimpulan")
    if kesimpulan_idx is None or kesimpulan_idx + 1 >= len(paragraphs):
        return

    end_idx = len(paragraphs)
    for idx in range(kesimpulan_idx + 1, len(paragraphs)):
        text = (paragraphs[idx].text or "").strip()
        if re.match(r"^[A-Z]\.", text):
            end_idx = idx
            break

    for idx in range(end_idx - 1, kesimpulan_idx, -1):
        _delete_paragraph(paragraphs[idx])

    anchor = paragraphs[kesimpulan_idx]
    current = anchor
    intro_text, numbered_items = _split_intro_and_numbered_items(kesimpulan_text)

    if intro_text:
        current = _insert_paragraph_after(current, intro_text)
        current.style = doc.styles["Normal"]
        _set_tight_paragraph_format(current)

    if numbered_items:
        for idx, item_text in enumerate(numbered_items, start=1):
            body = re.sub(r"^\d+\.\s*", "", item_text).strip()
            current = _insert_numbered_item_after(current, idx, body)
            current.style = doc.styles["Normal"]
            _set_numbered_paragraph_format(current, left_inch=0.5, hanging_inch=0.5)


def apply_title_model(doc, model_value):
    title_text = "Laporan Hasil Akuisisi Barang Bukti Elektronik"
    model_text = (model_value or "").strip()
    if model_text:
        title_text = f"{title_text} {model_text}"
    for paragraph in doc.paragraphs:
        if title_text in (paragraph.text or ""):
            return
        if "Laporan Hasil Akuisisi Barang Bukti Elektronik" in (paragraph.text or ""):
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = title_text
            else:
                paragraph.add_run(title_text)
            return


def append_ocr_section(doc, row):
    summary = (row["ocr_summary_text"] or "").strip()
    if not summary:
        return

    if summary.startswith("{"):
        return

    doc.add_heading("H. Data yang Ditemukan dari OCR", level=2)
    for line in summary.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.endswith(":"):
            doc.add_paragraph(text).runs[0].bold = True
        else:
            doc.add_paragraph(text)


def parse_dasar_items(form):
    # form may include per-row nomor and tanggal: dasar_nomor, dasar_tanggal
    jenis_list = form.getlist("dasar_jenis")
    isi_list = form.getlist("dasar_isi")
    nomor_list = form.getlist("dasar_nomor")
    tgl_list = form.getlist("dasar_tanggal")
    items = []
    for idx, jenis in enumerate(jenis_list):
        jenis_value = (jenis or "").strip()
        isi_value = (isi_list[idx] if idx < len(isi_list) else "" or "").strip()
        nomor_value = (nomor_list[idx] if idx < len(nomor_list) else "" or "").strip()
        tgl_value = (tgl_list[idx] if idx < len(tgl_list) else "" or "").strip()
        # Keep the user's typed text as the primary content so each row
        # remains an independent numbered item in the generated report.
        # Only fall back to a compact composed text when the main field is empty.
        if isi_value:
            isi_final = ' '.join(isi_value.split())
        else:
            parts = []
            if jenis_value:
                parts.append(jenis_value)
            if nomor_value:
                parts.append(f"No. {nomor_value}")
            if tgl_value:
                parts.append(f"tanggal {format_tanggal(tgl_value)}")
            isi_final = ' '.join([p for p in parts if p]).strip()
        if not isi_final:
            continue
        items.append({
            "jenis": jenis_value,
            "isi": isi_final,
            "nomor": nomor_value,
            "tanggal": tgl_value,
        })
    return items


def strip_indexing_from_text(text):
    if not text:
        return text
    lines = str(text).splitlines()
    out_lines = []
    skip_mode = False
    # common markers or keys that indicate indexing summary
    index_markers = ["--- Indeks", "--- Ringkasan Indexing", "Indeks:", "Ringkasan Indexing:"]
    index_keys = ["contacts", "chats", "images", "text", "installed applications", "applications", "all files", "archives", "databases", "videos", "instant messages", "user accounts"]
    for ln in lines:
        s = ln.strip()
        if any(m.lower() in s.lower() for m in index_markers):
            skip_mode = True
            continue
        if skip_mode:
            # stop skipping when reach an empty line
            if s == "":
                skip_mode = False
                continue
            # skip lines that look like indexing key/value pairs
            lower = s.lower()
            if any(lower.startswith(k) or (k + ':') in lower for k in index_keys):
                continue
            # if not an index-like line, stop skipping and include it
            skip_mode = False
        # also strip inline occurrences like '--- Indeks: ...' within a line
        if any(m.lower() in s.lower() for m in index_markers):
            continue
        out_lines.append(ln)
    return "\n".join(out_lines).strip()


def save_photos(files):
    paths = []
    for file in files:
        if not file or not file.filename:
            continue
        if not allowed_file(file.filename):
            continue
        safe_name = secure_filename(file.filename)
        final_name = f"{uuid.uuid4().hex}_{safe_name}"
        file_path = os.path.join(UPLOAD_DIR, final_name)
        file.save(file_path)
        paths.append(final_name)
    return paths


def format_pelaksanaan_akuisisi_section(doc, row):
    paragraphs = doc.paragraphs

    def find_heading_index(keyword):
        for idx, paragraph in enumerate(paragraphs):
            if keyword.lower() in (paragraph.text or "").lower():
                return idx
        return None

    g_idx = find_heading_index("G. Pelaksanaan Akuisisi")
    h_idx = find_heading_index("H. Data yang Ditemukan dari OCR")
    i_idx = find_heading_index("I. Kesimpulan")

    insert_before_idx = g_idx
    if insert_before_idx is None:
        insert_before_idx = h_idx if h_idx is not None else i_idx
    if insert_before_idx is None:
        return

    field_map = {
        'waktu dimulai': row['waktu_dimulai'] or '-',
        'waktu selesai': row['waktu_selesai'] or '-',
        'tempat': row['tempat_pelaksanaan'] or '-',
        'perangkat forensik digital yang digunakan': row['perangkat_forensik_digital'] or '-',
        'metode akuisisi': row['metode_akuisisi'] or '-',
        'status': row['status_akuisisi'] or '-',
        'nama file': row['nama_file_hash'] or '-',
        'nilai hash': row['nilai_hash'] or '-',
    }

    def clean_text(value):
        return re.sub(r'\s+', ' ', (value or '').strip()).lower()

    for table in doc.tables:
        for table_row in table.rows:
            if len(table_row.cells) < 2:
                continue
            label_cell = table_row.cells[0]
            value_cell = table_row.cells[1]
            label_text = clean_text(label_cell.text)
            for key, value in field_map.items():
                if key in label_text:
                    for paragraph in value_cell.paragraphs:
                        _delete_paragraph(paragraph)
                    value_cell.text = ''
                    new_paragraph = value_cell.add_paragraph(str(value))
                    new_paragraph.style = doc.styles['Normal']
                    _set_tight_paragraph_format(new_paragraph)
                    break

    if g_idx is None:
        anchor = paragraphs[insert_before_idx]
        heading = _insert_paragraph_before(anchor, "G. Pelaksanaan Akuisisi")
        heading.style = doc.styles["Heading 2"]
        current = heading
    else:
        current = paragraphs[g_idx]
        if g_idx + 1 < len(paragraphs):
            next_text = (paragraphs[g_idx + 1].text or "").strip()
            if not next_text:
                _delete_paragraph(paragraphs[g_idx + 1])

    field_rows = [
        ("Waktu dimulai", row["waktu_dimulai"] or "-"),
        ("Waktu selesai", row["waktu_selesai"] or "-"),
        ("Tempat", row["tempat_pelaksanaan"] or "-"),
        ("Perangkat forensik digital yang digunakan", row["perangkat_forensik_digital"] or "-"),
        ("Metode Akuisisi", row["metode_akuisisi"] or "-"),
        ("Status", row["status_akuisisi"] or "-"),
        ("Nama file", row["nama_file_hash"] or "-"),
        ("Nilai Hash", row["nilai_hash"] or "-"),
    ]

    for label, value in field_rows:
        if current is None:
            break
        current = _insert_paragraph_after(current, f"{label}: {value}")
        current.style = doc.styles["Normal"]
        _set_tight_paragraph_format(current)


def _get_file_lists_from_form(form, files):
    return {
        'foto_bb_paths': save_photos(files.getlist('fotos_bb')),
        'foto_hash_paths': save_photos(files.getlist('fotos_hash')),
        'foto_ekstraksi_paths': save_photos(files.getlist('fotos_ekstraksi')),
        'foto_indexing_paths': save_photos(files.getlist('fotos_indexing')),
    }


def _get_first_surat_dinas_data(form):
    jenis_list = form.getlist('dasar_jenis')
    nomor_list = form.getlist('dasar_nomor')
    tgl_list = form.getlist('dasar_tanggal')
    nomor_surat_col = ""
    tanggal_surat_col = ""
    for idx, jenis in enumerate(jenis_list):
        if (jenis or '').strip() == 'Surat Dinas':
            if idx < len(nomor_list):
                nomor_surat_col = (nomor_list[idx] or '').strip()
            if idx < len(tgl_list):
                tanggal_surat_col = (tgl_list[idx] or '').strip()
            break
    return nomor_surat_col, tanggal_surat_col


def _build_laporan_values(form, dasar_items, foto_bb_paths, foto_hash_paths, foto_ekstraksi_paths, foto_indexing_paths, nomor_surat_col, tanggal_surat_col, prosedur_value, status):
    pelaksanaan_values = build_pelaksanaan_akuisisi_values(form)
    ocr_fields = build_ocr_data_from_form(form)
    perangkat_info = build_perangkat_info_from_form(form)
    return (
        form.get('nama_instansi', '').strip(),
        form.get('alamat_instansi', '').strip(),
        nomor_surat_col,
        tanggal_surat_col,
        form.get('tanggal_serah_terima', '').strip(),
        form.get('komoditi', '').strip(),
        form.get('jenis_pelanggaran', '').strip(),
        form.get('jenis_barang_bukti', '').strip(),
        form.get('merek_model', '').strip(),
        (', '.join([v.strip() for v in form.getlist('imei') if v and v.strip()]) or ''),
        strip_indexing_from_text(form.get('keterangan_barang', '').strip()),
        form.get('ocr_raw', '').strip(),
        json.dumps(ocr_fields, ensure_ascii=False),
        build_ocr_summary_text_from_fields(ocr_fields),
        json.dumps(perangkat_info, ensure_ascii=False),
        json.dumps(dasar_items, ensure_ascii=True),
        json.dumps(foto_bb_paths, ensure_ascii=True),
        json.dumps(foto_hash_paths, ensure_ascii=True),
        json.dumps(foto_ekstraksi_paths, ensure_ascii=True),
        json.dumps(foto_indexing_paths, ensure_ascii=True),
        pelaksanaan_values['waktu_dimulai'],
        pelaksanaan_values['waktu_selesai'],
        pelaksanaan_values['tempat_pelaksanaan'],
        pelaksanaan_values['perangkat_forensik_digital'],
        pelaksanaan_values['metode_akuisisi'],
        pelaksanaan_values['status_akuisisi'],
        pelaksanaan_values['nama_file_hash'],
        pelaksanaan_values['nilai_hash'],
        form.get('maksud_tujuan', '').strip(),
        prosedur_value,
        form.get('kesimpulan', '').strip(),
        status,
        datetime.now().isoformat(timespec='seconds'),
    )


def _generate_nomor_laporan(conn, tanggal_laporan, jenis_field, provided_no_peralatan=''):
    jenis_map = {'Handphone': 'H', 'PC': 'PC', 'Laptop': 'LP'}
    jenis_code = jenis_map.get(jenis_field, (jenis_field[:2].upper() if jenis_field else 'H'))
    cur = conn.cursor()
    cur.execute('BEGIN IMMEDIATE')
    cur.execute('CREATE TABLE IF NOT EXISTS counters (name TEXT PRIMARY KEY, value INTEGER)')

    if provided_no_peralatan:
        try:
            no_peralatan_val = int(provided_no_peralatan)
        except Exception:
            no_peralatan_val = 0
    else:
        key_equip = f'equipseq:{jenis_code}'
        cur.execute('INSERT OR IGNORE INTO counters (name, value) VALUES (?, 0)', (key_equip,))
        cur.execute('SELECT value FROM counters WHERE name = ?', (key_equip,))
        row = cur.fetchone()
        cur_val = int(row['value']) if row and row['value'] is not None else 0
        no_peralatan_val = cur_val + 1
        cur.execute('UPDATE counters SET value = ? WHERE name = ?', (no_peralatan_val, key_equip))

    dt = _parse_laporan_datetime(tanggal_laporan)
    key_case = f'caseseq:{jenis_code}'
    case_next = _next_counter_value(cur, key_case)
    nomor_laporan_generated = _compose_nomor_laporan(jenis_code, dt, case_next)
    return cur, nomor_laporan_generated, jenis_code, key_case, dt


def _parse_laporan_datetime(tanggal_laporan):
    try:
        return datetime.fromisoformat(tanggal_laporan)
    except Exception:
        try:
            return datetime.strptime(tanggal_laporan, '%Y-%m-%d')
        except Exception:
            return datetime.now()


def _next_counter_value(cur, counter_name):
    cur.execute('INSERT OR IGNORE INTO counters (name, value) VALUES (?, 0)', (counter_name,))
    cur.execute('SELECT value FROM counters WHERE name = ?', (counter_name,))
    row = cur.fetchone()
    current_value = int(row['value']) if row and row['value'] is not None else 0
    next_value = current_value + 1
    cur.execute('UPDATE counters SET value = ? WHERE name = ?', (next_value, counter_name))
    return next_value


def _compose_nomor_laporan(jenis_code, dt, sequence_value):
    dd = f'{dt.day:02d}'
    mm = f'{dt.month:02d}'
    y2 = str(dt.year)[-2:]
    case_str = f'{sequence_value:02d}'
    return f'LDFOM.AK.{jenis_code}.{dd}.{mm}.{y2}.{case_str}'


def _insert_laporan_row(conn, nomor_laporan_generated, tanggal_laporan, form, dasar_items, foto_bb_paths, foto_hash_paths, foto_ekstraksi_paths, foto_indexing_paths, nomor_surat_col, tanggal_surat_col, prosedur_value, status):
    values = _build_laporan_values(
        form,
        dasar_items,
        foto_bb_paths,
        foto_hash_paths,
        foto_ekstraksi_paths,
        foto_indexing_paths,
        nomor_surat_col,
        tanggal_surat_col,
        prosedur_value,
        status,
    )
    conn.execute(
        """
        INSERT INTO laporan_akuisisi (
            nomor_laporan, tanggal_laporan, nama_instansi, alamat_instansi,
            nomor_surat_dinas, tanggal_surat_dinas, tanggal_serah_terima, komoditi, jenis_pelanggaran, jenis_barang_bukti,
            merek_model, imei, keterangan_barang, ocr_raw_text, ocr_fields_json, ocr_summary_text, perangkat_info_json, dasar_items, foto_paths,
            hash_foto_paths, ekstraksi_foto_paths, indexing_foto_paths,
            waktu_dimulai, waktu_selesai, tempat_pelaksanaan, perangkat_forensik_digital, metode_akuisisi, status_akuisisi, nama_file_hash, nilai_hash,
            maksud_tujuan, prosedur_akuisisi, kesimpulan, status, created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (nomor_laporan_generated, tanggal_laporan, *values),
    )


def _finalize_generated_docx(output_path, row, dasar_items, foto_bb_paths, hash_foto_paths, indexing_foto_paths):
    final_doc = Document(output_path)
    apply_title_model(final_doc, row['merek_model'])
    fill_template_hints(final_doc, row, foto_bb_paths, hash_foto_paths, indexing_foto_paths)
    format_pelaksanaan_akuisisi_section(final_doc, row)
    format_numbered_sections(final_doc, dasar_items, row['prosedur_akuisisi'] or '')
    align_dasar_indent_to_prosedur(final_doc)
    _fill_perangkat_info_table(final_doc, row)
    _fill_ocr_summary_table(final_doc, row)
    format_kesimpulan_section(final_doc, row['kesimpulan'] or '')
    append_ocr_section(final_doc, row)
    final_doc.save(output_path)
    normalize_docx_font(output_path)
    return output_path


def build_ocr_summary_text(fields, raw_text):
    if isinstance(fields, dict) and fields:
        return build_ocr_summary_text_from_fields(fields)
    return build_ocr_summary_text_from_fields(_parse_ocr_summary_stats(raw_text or ""))


def preprocess_image_for_ocr(path):
    # Pillow-only preprocessing to avoid heavy native dependencies.
    if not PIL_AVAILABLE:
        return None
    img = Image.open(path)
    if img.mode not in ('L', 'RGB'):
        img = img.convert('RGB')
    # upscale to improve OCR for small text
    w, h = img.size
    img = img.resize((int(w * 1.8), int(h * 1.8)))
    gray = img.convert('L')
    # simple threshold to improve contrast
    gray = gray.point(lambda p: 255 if p > 180 else 0)
    return gray


def _ocr_with_pytesseract_variants(path):
    if not PIL_AVAILABLE or not TESSERACT_AVAILABLE:
        return ''

    variants = []
    try:
        variants.append(Image.open(path))
    except Exception:
        pass

    try:
        original = Image.open(path)
        if original.mode not in ('L', 'RGB'):
            original = original.convert('RGB')
        variants.append(original)
    except Exception:
        pass

    processed = preprocess_image_for_ocr(path)
    if processed is not None:
        variants.append(processed)

    best_text = ''
    configs = ['--oem 3 --psm 6', '--oem 3 --psm 11', '--oem 3 --psm 4', '']
    langs = ['ind+eng', 'eng', None]

    for image in variants:
        for lang in langs:
            for config in configs:
                try:
                    kwargs = {}
                    if lang:
                        kwargs['lang'] = lang
                    if config:
                        kwargs['config'] = config
                    text = pytesseract.image_to_string(image, **kwargs)
                except Exception:
                    continue
                if text and len(text.strip()) > len(best_text.strip()):
                    best_text = text
                if best_text.strip() and len(best_text.strip()) > 50:
                    return best_text
    return best_text


@app.route("/")
def index():
    return redirect(url_for("dashboard"))


@app.route("/dashboard")
def dashboard():
    conn = get_conn()
    
    # Total laporan
    total = conn.execute("SELECT COUNT(*) as cnt FROM laporan_akuisisi").fetchone()['cnt']
    
    # Draft dan Final
    draft = conn.execute("SELECT COUNT(*) as cnt FROM laporan_akuisisi WHERE status = 'draft'").fetchone()['cnt']
    final = conn.execute("SELECT COUNT(*) as cnt FROM laporan_akuisisi WHERE status = 'final'").fetchone()['cnt']
    
    # Jenis barang bukti unik
    jenis = conn.execute("SELECT COUNT(DISTINCT jenis_barang_bukti) as cnt FROM laporan_akuisisi WHERE jenis_barang_bukti IS NOT NULL AND jenis_barang_bukti != ''").fetchone()['cnt']
    
    # 5 laporan terbaru (berdasarkan id/waktu)
    recent = conn.execute(
        "SELECT id, nomor_laporan, tanggal_laporan, nama_instansi, merek_model, status FROM laporan_akuisisi ORDER BY id DESC LIMIT 5"
    ).fetchall()
    
    conn.close()
    
    return render_template(
        "dashboard.html",
        total_laporan=total,
        draft_count=draft,
        final_count=final,
        jenis_count=jenis,
        recent_laporan=recent
    )


@app.route("/api/analytics/chart-data")
def analytics_chart_data():
    """Return chart data for dashboard visualization"""
    conn = get_conn()
    
    # Status breakdown
    status_data = conn.execute(
        "SELECT status, COUNT(*) as count FROM laporan_akuisisi GROUP BY status"
    ).fetchall()
    
    # Jenis barang bukti breakdown
    jenis_data = conn.execute(
        "SELECT jenis_barang_bukti, COUNT(*) as count FROM laporan_akuisisi WHERE jenis_barang_bukti IS NOT NULL AND jenis_barang_bukti != '' GROUP BY jenis_barang_bukti"
    ).fetchall()
    
    # Laporan per bulan (last 12 months)
    monthly_data = conn.execute(
        """SELECT strftime('%Y-%m', tanggal_laporan) as month, COUNT(*) as count 
           FROM laporan_akuisisi 
           GROUP BY month 
           ORDER BY month DESC LIMIT 12"""
    ).fetchall()
    
    # Top instansi
    top_instansi = conn.execute(
        "SELECT nama_instansi, COUNT(*) as count FROM laporan_akuisisi WHERE nama_instansi IS NOT NULL AND nama_instansi != '' GROUP BY nama_instansi ORDER BY count DESC LIMIT 5"
    ).fetchall()
    
    conn.close()
    
    return jsonify({
        'status': [{'status': s['status'], 'count': s['count']} for s in status_data],
        'jenis': [{'jenis': j['jenis_barang_bukti'], 'count': j['count']} for j in jenis_data],
        'monthly': [{'month': m['month'], 'count': m['count']} for m in reversed(list(monthly_data))],
        'top_instansi': [{'instansi': i['nama_instansi'], 'count': i['count']} for i in top_instansi],
    })


@app.route("/api/export/csv")
def export_csv():
    """Export all laporan as CSV"""
    import csv as csv_module
    from io import StringIO
    
    conn = get_conn()
    laporan = conn.execute("SELECT * FROM laporan_akuisisi ORDER BY id DESC").fetchall()
    conn.close()
    
    if not laporan:
        flash('Tidak ada data untuk diekspor.', 'warning')
        return redirect(url_for('dashboard'))
    
    # Create CSV in memory
    output = StringIO()
    writer = csv_module.DictWriter(output, fieldnames=['id', 'nomor_laporan', 'tanggal_laporan', 'nama_instansi', 'komoditi', 'jenis_barang_bukti', 'merek_model', 'imei', 'status', 'created_at'])
    writer.writeheader()
    
    for row in laporan:
        writer.writerow({
            'id': row['id'],
            'nomor_laporan': row['nomor_laporan'],
            'tanggal_laporan': row['tanggal_laporan'],
            'nama_instansi': row['nama_instansi'],
            'komoditi': row['komoditi'],
            'jenis_barang_bukti': row['jenis_barang_bukti'],
            'merek_model': row['merek_model'],
            'imei': row['imei'],
            'status': row['status'],
            'created_at': row['created_at'],
        })
    
    from flask import Response
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=laporan_akuisisi.csv'}
    )


@app.route("/laporan")
def list_laporan():
    conn = get_conn()
    laporan = conn.execute("SELECT * FROM laporan_akuisisi ORDER BY id ASC").fetchall()
    conn.close()
    return render_template("list.html", laporan=laporan)


@app.route("/laporan/new", methods=["GET", "POST"])
def create_laporan():
    if request.method == "POST":
        tanggal_laporan = request.form.get("tanggal_laporan", "").strip()
        status = request.form.get("status", "draft").strip()

        if not tanggal_laporan:
            flash("Tanggal laporan wajib diisi.", "danger")
            return redirect(url_for("create_laporan"))

        dasar_items = parse_dasar_items(request.form)
        nomor_surat_col, tanggal_surat_col = _get_first_surat_dinas_data(request.form)
        foto_paths = _get_file_lists_from_form(request.form, request.files)
        foto_bb_paths = foto_paths['foto_bb_paths']
        foto_hash_paths = foto_paths['foto_hash_paths']
        foto_ekstraksi_paths = foto_paths['foto_ekstraksi_paths']
        foto_indexing_paths = foto_paths['foto_indexing_paths']
        prosedur_value = build_prosedur_akuisisi(
            request.form.get("prosedur_option"),
            request.form.get("prosedur_akuisisi", ""),
        )

        conn = get_conn()
        try:
            cur, nomor_laporan_generated, jenis_code, key_case, dt = _generate_nomor_laporan(
                conn,
                tanggal_laporan,
                request.form.get("jenis_barang_bukti", "Handphone").strip(),
                request.form.get('no_peralatan', '').strip(),
            )
            try:
                _insert_laporan_row(
                    conn,
                    nomor_laporan_generated,
                    tanggal_laporan,
                    request.form,
                    dasar_items,
                    foto_bb_paths,
                    foto_hash_paths,
                    foto_ekstraksi_paths,
                    foto_indexing_paths,
                    nomor_surat_col,
                    tanggal_surat_col,
                    prosedur_value,
                    status,
                )
                conn.commit()
            except sqlite3.IntegrityError:
                try:
                    new_case = _next_counter_value(cur, key_case)
                    nomor_laporan_generated = _compose_nomor_laporan(jenis_code, dt, new_case)
                    _insert_laporan_row(
                        conn,
                        nomor_laporan_generated,
                        tanggal_laporan,
                        request.form,
                        dasar_items,
                        foto_bb_paths,
                        foto_hash_paths,
                        foto_ekstraksi_paths,
                        foto_indexing_paths,
                        nomor_surat_col,
                        tanggal_surat_col,
                        prosedur_value,
                        status,
                    )
                    conn.commit()
                except Exception:
                    conn.rollback()
                    flash("Gagal menyimpan laporan (konflik nomor). Coba lagi.", "danger")
                    return redirect(url_for("create_laporan"))
            flash("Laporan berhasil disimpan.", "success")
            return redirect(url_for("list_laporan"))
        except Exception as e:
            try:
                conn.rollback()
            except:
                pass
            flash(f"Gagal menyimpan laporan: {e}", "danger")
            return redirect(url_for("create_laporan"))
        finally:
            conn.close()
    instansi_list = load_instansi()
    return render_template("form.html", instansi_list=instansi_list, ocr_fields={}, perangkat_info={})


@app.route('/api/instansi')
def api_instansi():
    # Return current instansi list as JSON so client can fetch fresh data
    inst = load_instansi()
    return jsonify(inst)


@app.route('/api/surats')
def api_surats():
    """Return predefined surat data (Surat Dinas and Surat Tugas) as JSON."""
    return jsonify(PREDEFINED_SURATS)


@app.route('/api/ocr', methods=['POST'])
def api_ocr():
    """OCR sebuah gambar dan kembalikan teks + nilai hash terparse.

    Menerima salah satu dari: file baru pada field 'file', ATAU nama file yang
    sudah tersimpan sebelumnya pada field 'filename' (untuk mode edit)."""
    f = request.files.get('file')
    if f:
        if not allowed_file(f.filename):
            return jsonify({'error': 'file type not allowed'}), 400
        safe_name = secure_filename(f.filename)
        fname = f"{uuid.uuid4().hex}_{safe_name}"
        out_path = os.path.join(UPLOAD_DIR, fname)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        f.save(out_path)
    else:
        # File yang sudah tersimpan (mode edit): validasi agar tetap di UPLOAD_DIR.
        requested = (request.form.get('filename') or '').strip()
        fname = secure_filename(requested)
        if not fname:
            return jsonify({'error': 'no file uploaded'}), 400
        out_path = os.path.join(UPLOAD_DIR, fname)
        if os.path.dirname(os.path.abspath(out_path)) != os.path.abspath(UPLOAD_DIR) or not os.path.exists(out_path):
            return jsonify({'error': 'file not found'}), 404

    # Check dependencies
    missing = []
    if not PIL_AVAILABLE:
        missing.append('Pillow')
    if not WINRT_AVAILABLE:
        missing.append('winsdk')
    if not TESSERACT_AVAILABLE:
        missing.append('pytesseract')
    if not _USE_GOOGLE_VISION and not WINRT_AVAILABLE and not TESSERACT_AVAILABLE:
        return jsonify({
            'error': 'no OCR backend available',
            'missing': missing,
            'hint': 'Install winsdk for Windows OCR or configure Google Vision'
        }), 500

    text = ''
    # Prefer Google Vision if available and credentials set
    if _USE_GOOGLE_VISION:
        try:
            text = ocr_with_google(out_path)
        except Exception:
            text = ''

    # Try Windows OCR before pytesseract if Google Vision is not used or failed
    if not text and WINRT_AVAILABLE:
        try:
            text = ocr_with_windows(out_path)
        except Exception:
            text = ''

    # Fallback to local pytesseract if Google Vision not used or failed
    if not text and TESSERACT_AVAILABLE:
        try:
            text = _ocr_with_pytesseract_variants(out_path)
        except Exception:
            text = ''

    # OCR khusus (grayscale/upscale) untuk pembacaan hex terbaik pada gambar
    # hash; fallback ke teks generik bila hasilnya kosong.
    hash_text = ''
    if TESSERACT_AVAILABLE:
        try:
            hash_text = _ocr_hash_image(out_path)
        except Exception:
            hash_text = ''
    hashes, hash_warnings = parse_hash_values(hash_text or text)
    return jsonify({
        'text': text,
        'fields': {},
        'hashes': hashes,
        'nilai_hash_text': format_nilai_hash_text(hashes),
        'warnings': hash_warnings,
        'saved_as': fname,
    })


@app.route('/api/reserve_nomor', methods=['POST', 'GET'])
def api_reserve_nomor():
    """Atomically reserve and return the next nomor laporan.
    Returns JSON: {nomor: str, seq: int}
    """
    # read parameters (operation: AN/AK, jenis: equipment code like H/PC/LP, no_peralatan optional, tanggal optional)
    data = request.get_json(force=False, silent=True) or request.form or {}
    operation = (data.get('operation') or data.get('operasi') or 'AK').upper()
    jenis = (data.get('jenis') or data.get('jenis_kode') or data.get('jenis_barang_bukti') or 'H')
    no_peralatan = data.get('no_peralatan')
    tanggal = data.get('tanggal') or data.get('tanggal_laporan')

    # normalize jenis (take first token or mapping)
    jenis = (jenis or '').strip()
    if not jenis:
        jenis = 'H'

    # determine date parts
    try:
        if tanggal:
            dt = datetime.fromisoformat(tanggal)
        else:
            dt = datetime.now()
    except Exception:
        try:
            dt = datetime.strptime(tanggal, '%Y-%m-%d')
        except Exception:
            dt = datetime.now()
    y2 = str(dt.year)[-2:]
    mm = f"{dt.month:02d}"

    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute('BEGIN IMMEDIATE')
        # ensure counters table
        cur.execute("CREATE TABLE IF NOT EXISTS counters (name TEXT PRIMARY KEY, value INTEGER)")

        # if no_peralatan not provided, increment a persistent per-jenis equipment counter
        if not no_peralatan:
            key_equip = f"equipseq:{jenis}"
            cur.execute("INSERT OR IGNORE INTO counters (name, value) VALUES (?, 0)", (key_equip,))
            cur.execute("SELECT value FROM counters WHERE name = ?", (key_equip,))
            row = cur.fetchone()
            cur_val = int(row['value']) if row and row['value'] is not None else 0
            next_equip = cur_val + 1
            cur.execute("UPDATE counters SET value = ? WHERE name = ?", (next_equip, key_equip))
            no_peralatan_val = next_equip
        else:
            try:
                no_peralatan_val = int(str(no_peralatan).strip())
            except Exception:
                no_peralatan_val = 0

        # document sequence per operation+jenis+YYYYMM
        key_doc = f"docseq:{operation}:{jenis}:{dt.year}{dt.month:02d}"
        cur.execute("INSERT OR IGNORE INTO counters (name, value) VALUES (?, 0)", (key_doc,))
        cur.execute("SELECT value FROM counters WHERE name = ?", (key_doc,))
        row = cur.fetchone()
        doc_cur = int(row['value']) if row and row['value'] is not None else 0
        next_doc = doc_cur + 1
        cur.execute("UPDATE counters SET value = ? WHERE name = ?", (next_doc, key_doc))

        conn.commit()

        # format components
        no_peralatan_str = f"{no_peralatan_val:02d}"
        docseq_str = f"{next_doc:02d}"
        nomor = f"LDFOM.{operation}.{jenis}.{no_peralatan_str}.{mm}.{y2}.{docseq_str}"
        return jsonify({"nomor": nomor, "no_peralatan": no_peralatan_val, "doc_seq": next_doc})
    except Exception as e:
        try:
            conn.rollback()
        except:
            pass
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route('/debug/instansi')
def debug_instansi():
    csv_path = os.path.join(BASE_DIR, 'data', 'instansi.csv')
    info = {'csv_path': csv_path, 'exists': os.path.exists(csv_path)}
    try:
        if os.path.exists(csv_path):
            stat = os.stat(csv_path)
            info['size'] = stat.st_size
            info['mtime'] = stat.st_mtime
            with open(csv_path, 'r', encoding='utf-8', newline='') as f:
                lines = f.readlines()
            info['line_count'] = len(lines)
            # Try to parse header and first 20 names
            try:
                f = open(csv_path, newline='', encoding='utf-8')
                reader = csv.DictReader(f)
                names = []
                for i, row in enumerate(reader):
                    if i >= 20:
                        break
                    name = row.get('Nama UPT') or row.get('Nama') or ''
                    names.append(name)
                f.close()
                info['sample_names_count'] = len(names)
                info['sample_names'] = names
            except Exception as e:
                info['parse_error'] = str(e)
        else:
            info['error'] = 'file not found'
    except Exception as e:
        info['error'] = str(e)
    return jsonify(info)


@app.route('/debug/instansi/content')
def debug_instansi_content():
    csv_path = os.path.join(BASE_DIR, 'data', 'instansi.csv')
    if not os.path.exists(csv_path):
        return jsonify({'error': 'file not found', 'csv_path': csv_path})
    try:
        with open(csv_path, 'rb') as f:
            raw = f.read()
        import hashlib

        sha256 = hashlib.sha256(raw).hexdigest()
        # return first 2000 chars of text for safety
        text = raw.decode('utf-8', errors='replace')
        snippet = text[:20000]
        return jsonify({'csv_path': csv_path, 'size': len(raw), 'sha256': sha256, 'snippet': snippet})
    except Exception as e:
        return jsonify({'error': str(e)})


@app.route("/laporan/<int:laporan_id>/generate-docx")
def generate_docx(laporan_id):
    if not os.path.exists(TEMPLATE_PATH):
        ensure_default_template()
        flash("Template default otomatis dibuat: template_laporan.docx", "info")

    conn = get_conn()
    row = conn.execute("SELECT * FROM laporan_akuisisi WHERE id = ?", (laporan_id,)).fetchone()
    conn.close()
    if not row:
        flash("Data laporan tidak ditemukan.", "danger")
        return redirect(url_for("list_laporan"))

    dasar_items = json.loads(row["dasar_items"] or "[]")
    foto_bb_paths = json.loads(row["foto_paths"] or "[]")
    hash_foto_paths = json.loads(row["hash_foto_paths"] or "[]")
    ekstraksi_foto_paths = json.loads(row["ekstraksi_foto_paths"] or "[]")
    indexing_foto_paths = json.loads(row["indexing_foto_paths"] or "[]")

    work_template = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    work_template.close()
    shutil.copy(TEMPLATE_PATH, work_template.name)
    repair_template_jinja_tags(work_template.name)

    doc = DocxTemplate(work_template.name)
    context = build_render_context(doc, row, dasar_items, foto_bb_paths)
    doc.render(context)

    output_name = f"laporan_{row['nomor_laporan'].replace('.', '_')}_{laporan_id}.docx"
    output_path = os.path.join(GENERATED_DIR, output_name)
    doc.save(output_path)
    os.unlink(work_template.name)
    _finalize_generated_docx(output_path, row, dasar_items, foto_bb_paths, hash_foto_paths, indexing_foto_paths)
    return send_from_directory(GENERATED_DIR, output_name, as_attachment=True)


@app.route('/laporan/<int:laporan_id>/edit', methods=['GET', 'POST'])
def edit_laporan(laporan_id):
    conn = get_conn()
    row = conn.execute('SELECT * FROM laporan_akuisisi WHERE id = ?', (laporan_id,)).fetchone()
    if not row:
        conn.close()
        flash('Laporan tidak ditemukan.', 'danger')
        return redirect(url_for('list_laporan'))

    # Only allow editing if the laporan was created as draft
    if (row['status'] or '').lower() != 'draft':
        conn.close()
        flash('Hanya laporan dengan status draft yang bisa diedit.', 'warning')
        return redirect(url_for('list_laporan'))

    if request.method == 'POST':
        # update fields; nomor_laporan can now be edited
        nomor_laporan_baru = request.form.get('nomor_laporan', '').strip()
        tanggal_laporan = request.form.get('tanggal_laporan', '').strip()
        status = request.form.get('status', 'draft').strip()
        
        # validate nomor_laporan is not empty and is unique (if changed)
        if not nomor_laporan_baru:
            conn.close()
            flash('Nomor laporan wajib diisi.', 'danger')
            return redirect(url_for('edit_laporan', laporan_id=laporan_id))
        
        # check if nomor_laporan changed and if new one already exists
        if nomor_laporan_baru != row['nomor_laporan']:
            existing = conn.execute('SELECT id FROM laporan_akuisisi WHERE nomor_laporan = ?', (nomor_laporan_baru,)).fetchone()
            if existing:
                conn.close()
                flash('Nomor laporan sudah digunakan. Pilih nomor yang berbeda.', 'danger')
                return redirect(url_for('edit_laporan', laporan_id=laporan_id))
        
        dasar_items = parse_dasar_items(request.form)

        # determine first Surat Dinas (if any) to populate nomor_surat_dinas / tanggal_surat_dinas columns
        jenis_list = request.form.getlist('dasar_jenis')
        nomor_list = request.form.getlist('dasar_nomor')
        tgl_list = request.form.getlist('dasar_tanggal')
        nomor_surat_col = ""
        tanggal_surat_col = ""
        for idx, j in enumerate(jenis_list):
            if (j or '').strip() == 'Surat Dinas':
                if idx < len(nomor_list):
                    nomor_surat_col = (nomor_list[idx] or '').strip()
                if idx < len(tgl_list):
                    tanggal_surat_col = (tgl_list[idx] or '').strip()
                break

        # handle photos: append new uploads to existing lists
        foto_bb_paths_existing = json.loads(row['foto_paths'] or '[]')
        foto_hash_paths_existing = json.loads(row['hash_foto_paths'] or '[]')
        foto_ekstraksi_existing = json.loads(row['ekstraksi_foto_paths'] or '[]')
        foto_indexing_existing = json.loads(row['indexing_foto_paths'] or '[]')
        # handle removals requested by user during edit
        remove_bb = request.form.getlist('remove_foto_bb')
        remove_hash = request.form.getlist('remove_foto_hash')
        remove_ekstraksi = request.form.getlist('remove_foto_ekstraksi')
        remove_indexing = request.form.getlist('remove_foto_indexing')
        # helper to delete files and remove from list
        def _remove_files(list_paths, removals):
            kept = []
            for fname in list_paths:
                if fname in removals:
                    try:
                        full = os.path.join(UPLOAD_DIR, fname)
                        if os.path.exists(full):
                            os.remove(full)
                    except Exception:
                        pass
                else:
                    kept.append(fname)
            return kept

        foto_bb_paths_existing = _remove_files(foto_bb_paths_existing, set(remove_bb))
        foto_hash_paths_existing = _remove_files(foto_hash_paths_existing, set(remove_hash))
        foto_ekstraksi_existing = _remove_files(foto_ekstraksi_existing, set(remove_ekstraksi))
        foto_indexing_existing = _remove_files(foto_indexing_existing, set(remove_indexing))
        new_bb = save_photos(request.files.getlist('fotos_bb'))
        new_hash = save_photos(request.files.getlist('fotos_hash'))
        new_ekstraksi = save_photos(request.files.getlist('fotos_ekstraksi'))
        new_indexing = save_photos(request.files.getlist('fotos_indexing'))

        foto_bb_paths = foto_bb_paths_existing + new_bb
        foto_hash_paths = foto_hash_paths_existing + new_hash
        foto_ekstraksi_paths = foto_ekstraksi_existing + new_ekstraksi
        foto_indexing_paths = foto_indexing_existing + new_indexing

        ocr_fields = build_ocr_data_from_form(request.form)
        ocr_fields_json = json.dumps(ocr_fields, ensure_ascii=False)
        ocr_summary_text = build_ocr_summary_text_from_fields(ocr_fields)
        perangkat_info = build_perangkat_info_from_form(request.form)
        perangkat_info_json = json.dumps(perangkat_info, ensure_ascii=False)

        try:
            # debug: print submitted keys to help diagnose missing updates
            try:
                print(f"DEBUG edit_laporan submit id={laporan_id} keys={list(request.form.keys())}")
            except Exception:
                pass
            conn.execute(
                                """
                                UPDATE laporan_akuisisi SET
                                    nomor_laporan = ?, tanggal_laporan = ?, nama_instansi = ?, alamat_instansi = ?,
                                    tanggal_serah_terima = ?, komoditi = ?, jenis_pelanggaran = ?, jenis_barang_bukti = ?,
                                    merek_model = ?, imei = ?, keterangan_barang = ?, ocr_raw_text = ?, ocr_fields_json = ?, ocr_summary_text = ?, perangkat_info_json = ?, dasar_items = ?, foto_paths = ?,
                                    hash_foto_paths = ?, ekstraksi_foto_paths = ?, indexing_foto_paths = ?,
                                    waktu_dimulai = ?, waktu_selesai = ?, tempat_pelaksanaan = ?, perangkat_forensik_digital = ?, metode_akuisisi = ?, status_akuisisi = ?, nama_file_hash = ?, nilai_hash = ?,
                                    maksud_tujuan = ?, prosedur_akuisisi = ?, kesimpulan = ?, status = ?,
                                    nomor_surat_dinas = ?, tanggal_surat_dinas = ?
                                WHERE id = ?
                                """,
                                (
                                        nomor_laporan_baru,
                                        tanggal_laporan,
                                        request.form.get('nama_instansi', '').strip(),
                                        request.form.get('alamat_instansi', '').strip(),
                                        request.form.get('tanggal_serah_terima', '').strip(),
                                        request.form.get('komoditi', '').strip(),
                                        request.form.get('jenis_pelanggaran', '').strip(),
                                        request.form.get('jenis_barang_bukti', '').strip(),
                                        request.form.get('merek_model', '').strip(),
                                        (', '.join([v.strip() for v in request.form.getlist('imei') if v and v.strip()]) or ''),
                                        strip_indexing_from_text(request.form.get('keterangan_barang', '').strip()),
                                        request.form.get('ocr_raw', '').strip(),
                                        ocr_fields_json,
                                        ocr_summary_text,
                                        perangkat_info_json,
                                        json.dumps(dasar_items, ensure_ascii=True),
                                        json.dumps(foto_bb_paths, ensure_ascii=True),
                                        json.dumps(foto_hash_paths, ensure_ascii=True),
                                        json.dumps(foto_ekstraksi_paths, ensure_ascii=True),
                                        json.dumps(foto_indexing_paths, ensure_ascii=True),
                                        request.form.get('waktu_dimulai', '').strip(),
                                        request.form.get('waktu_selesai', '').strip(),
                                        request.form.get('tempat_pelaksanaan', '').strip(),
                                        request.form.get('perangkat_forensik_digital', '').strip(),
                                        request.form.get('metode_akuisisi', '').strip(),
                                        request.form.get('status_akuisisi', '').strip(),
                                        request.form.get('nama_file_hash', '').strip(),
                                        request.form.get('nilai_hash', '').strip(),
                                        request.form.get('maksud_tujuan', '').strip(),
                                        request.form.get('prosedur_akuisisi', '').strip(),
                                        request.form.get('kesimpulan', '').strip(),
                                        status,
                                        nomor_surat_col,
                                        tanggal_surat_col,
                                        laporan_id,
                                ),
            )
            conn.commit()
            flash('Laporan berhasil diperbarui.', 'success')
            try:
                print(f"DEBUG: laporan {laporan_id} updated successfully")
            except Exception:
                pass
            return redirect(url_for('list_laporan'))
        except Exception as e:
            conn.rollback()
            flash(f'Gagal memperbarui laporan: {e}', 'danger')
            return redirect(url_for('edit_laporan', laporan_id=laporan_id))
        finally:
            conn.close()

    # GET: render form with existing data
    dasar_items = json.loads(row['dasar_items'] or '[]')
    foto_bb_paths = json.loads(row['foto_paths'] or '[]')
    foto_hash_paths = json.loads(row['hash_foto_paths'] or '[]')
    foto_ekstraksi_paths = json.loads(row['ekstraksi_foto_paths'] or '[]')
    foto_indexing_paths = json.loads(row['indexing_foto_paths'] or '[]')
    instansi_list = load_instansi()
    # sanitize keterangan for display (remove previously stored indexing lines)
    row_dict = dict(row)
    row_dict['keterangan_barang'] = strip_indexing_from_text(row_dict.get('keterangan_barang', ''))
    try:
        ocr_fields = json.loads(row_dict.get('ocr_fields_json') or '{}')
        if not isinstance(ocr_fields, dict):
            ocr_fields = {}
    except Exception:
        ocr_fields = _parse_ocr_summary_stats(row_dict.get('ocr_summary_text') or '')
    try:
        perangkat_info = json.loads(row_dict.get('perangkat_info_json') or '{}')
        if not isinstance(perangkat_info, dict):
            perangkat_info = {}
    except Exception:
        perangkat_info = {}
    conn.close()
    return render_template(
        'form.html',
        laporan=row_dict,
        dasar_items=dasar_items,
        instansi_list=instansi_list,
        foto_bb_paths=foto_bb_paths,
        foto_hash_paths=foto_hash_paths,
        foto_ekstraksi_paths=foto_ekstraksi_paths,
        foto_indexing_paths=foto_indexing_paths,
        ocr_fields=ocr_fields,
        perangkat_info=perangkat_info,
    )


@app.route("/uploads/<path:filename>")
def download_upload(filename):
    return send_from_directory(UPLOAD_DIR, filename, as_attachment=False)


if __name__ == "__main__":
    init_db()
    app.run(debug=True)
